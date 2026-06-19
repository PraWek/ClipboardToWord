import os
import re
import tempfile
import uuid
from flask import Flask, request, render_template, send_file
import pypandoc
import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

app = Flask(__name__)


def generate_reference_doc(settings):
    """Генерирует Word-шаблон на лету с расширенными метриками ГОСТ/БМСТУ."""
    doc = docx.Document()

    # 1. Настройка основного текста (стиль Normal)
    style_normal = doc.styles['Normal']
    style_normal.font.name = settings['font_name']
    style_normal.font.size = Pt(float(settings['font_size']))

    # Абзацные метрики основного текста
    pf = style_normal.paragraph_format
    pf.line_spacing = float(settings['line_spacing'])
    pf.first_line_indent = Cm(float(settings['first_line_indent']))

    alignments = {
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
    }
    pf.alignment = alignments.get(settings['alignment'], WD_ALIGN_PARAGRAPH.JUSTIFY)

    # 2. Настройка полей страницы
    for section in doc.sections:
        section.left_margin = Cm(float(settings['margin_left']))
        section.right_margin = Cm(float(settings['margin_right']))
        section.top_margin = Cm(float(settings['margin_top']))
        section.bottom_margin = Cm(float(settings['margin_bottom']))

    # 3. Настройка заголовков (чтобы не сбрасывались на стандартные)
    for i in range(1, 5):
        heading_style = f'Heading {i}'
        if heading_style in doc.styles:
            doc.styles[heading_style].font.name = settings['font_name']
            doc.styles[heading_style].font.color.rgb = docx.shared.RGBColor(0, 0, 0)  # Строгий черный по ГОСТу

    # 4. Настройка стиля для Листингов/Кода (Pandoc использует 'Source Code')
    try:
        style_code = doc.styles['Source Code']
    except KeyError:
        # Если стиля нет в базовом наборе python-docx, создаем его базовый вариант
        style_code = doc.styles.add_style('Source Code', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)

    style_code.font.name = 'Consolas'  # Моноширинный шрифт для кода по умолчанию
    style_code.font.size = Pt(11.0)
    style_code.paragraph_format.first_line_indent = Cm(0)  # У листингов нет красной строки
    style_code.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Без выравнивания по ширине

    # Сохранение временного файла
    temp_dir = tempfile.gettempdir()
    template_path = os.path.join(temp_dir, f"template_{uuid.uuid4().hex}.docx")
    doc.save(template_path)
    return template_path


def process_text_to_word(raw_text, output_path, settings):
    # Исправление синтаксиса формул (убираем пробелы внутри знаков доллара)
    clean_text = re.sub(r'\\\[\s*', '$$', raw_text)
    clean_text = re.sub(r'\s*\\\]', '$$', clean_text)
    clean_text = re.sub(r'\\\(\s*', '$', clean_text)
    clean_text = re.sub(r'\s*\\\)', '$', clean_text)

    template_path = generate_reference_doc(settings)
    extra_args = [f'--reference-doc={template_path}']

    try:
        pypandoc.convert_text(
            clean_text,
            'docx',
            format='markdown',
            outputfile=output_path,
            extra_args=extra_args
        )
    finally:
        if os.path.exists(template_path):
            os.remove(template_path)


@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        settings = {
            'font_name': request.form.get('font_name', 'Times New Roman'),
            'font_size': request.form.get('font_size', '14'),
            'line_spacing': request.form.get('line_spacing', '1.5'),
            'first_line_indent': request.form.get('first_line_indent', '1.25'),
            'alignment': request.form.get('alignment', 'justify'),
            'margin_left': request.form.get('margin_left', '3.0'),
            'margin_right': request.form.get('margin_right', '1.0'),
            'margin_top': request.form.get('margin_top', '2.0'),
            'margin_bottom': request.form.get('margin_bottom', '2.0')
        }

        raw_text = ""
        if 'file' in request.files and request.files['file'].filename != '':
            uploaded_file = request.files['file']
            ext = uploaded_file.filename.rsplit('.', 1)[-1].lower()
            try:
                if ext in ['txt', 'md']:
                    raw_text = uploaded_file.read().decode('utf-8')
                elif ext == 'docx':
                    doc = docx.Document(uploaded_file)
                    raw_text = "\n".join([p.text for p in doc.paragraphs])
                else:
                    return "Формат не поддерживается", 400
            except Exception as e:
                return f"Ошибка чтения файла: {e}", 500
        else:
            raw_text = request.form.get('text_input', '')

        if not raw_text.strip():
            return "Нет данных для обработки.", 400

        output_filename = "converted_result.docx"
        try:
            process_text_to_word(raw_text, output_filename, settings)
            return send_file(
                output_filename,
                as_attachment=True,
                download_name="Report_Result.docx"
            )
        except Exception as e:
            return f"Ошибка конвертации: {e}", 500

    return render_template('index.html')


if __name__ == "__main__":
    app.run(debug=True, host='0.0.0.0', port=5000)